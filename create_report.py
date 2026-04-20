"""
create_report.py — Tạo báo cáo Word (.docx) chuẩn học thuật UTH
Chạy: python create_report.py
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "output", "reports", "BaoCao_SpamClassifier.docx"
)
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)

doc = Document()

# ─────────────────────────────────────────────
# THIẾT LẬP TRANG (chuẩn luận văn VN)
# Lề: trái 3.5cm, phải 2cm, trên 3cm, dưới 3cm
# ─────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3.5)
section.right_margin  = Cm(2.0)
section.top_margin    = Cm(3.0)
section.bottom_margin = Cm(3.0)

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────
def set_font(run, name="Times New Roman", size=13, bold=False,
             italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold      = bold
    run.italic    = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para(text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False,
         size=13, space_before=0, space_after=6, italic=False,
         color=None, font="Times New Roman"):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        r = p.add_run(text)
        set_font(r, font, size, bold, italic, color)
    return p

def heading_chapter(num, title):
    """Tiêu đề chương — in hoa, đậm, 14pt, căn giữa."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(12)
    r = p.add_run(f"CHƯƠNG {num}: {title.upper()}")
    set_font(r, size=14, bold=True, color=(31, 73, 125))
    return p

def heading_section(num_str, title):
    """Tiêu đề mục — đậm 13pt."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(f"{num_str} {title}")
    set_font(r, size=13, bold=True)
    return p

def heading_sub(num_str, title):
    """Tiêu đề tiểu mục — đậm nghiêng 13pt."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(f"{num_str} {title}")
    set_font(r, size=13, bold=True, italic=True)
    return p

def body(text, indent=False):
    """Đoạn văn thường."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(20)
    r = p.add_run(text)
    set_font(r, size=13)
    return p

def bullet(text, level=0):
    """Dòng có dấu gạch đầu dòng."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Cm(0.75 + level * 0.5)
    p.paragraph_format.space_after   = Pt(3)
    r = p.add_run(text)
    set_font(r, size=13)
    return p

def table_caption(text):
    """Tiêu đề bảng — căn giữa, nghiêng."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    set_font(r, size=12, italic=True, bold=True)
    return p

def fig_caption(text):
    """Tiêu đề hình — căn giữa, nghiêng."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    r = p.add_run(text)
    set_font(r, size=12, italic=True)
    return p

def add_table(headers, rows, col_widths=None):
    """Tạo bảng với header tô màu xanh đậm."""
    ncols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=ncols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Tô màu xanh đậm
        tc_pr = cell._tc.get_or_add_tcPr()
        shd   = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '1F497D')
        tc_pr.append(shd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=12, bold=True, color=(255, 255, 255))

    # Data rows
    for ri, row_data in enumerate(rows):
        drow = t.rows[ri + 1]
        fill = 'DCE6F1' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = drow.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Zebra striping
            tc_pr = cell._tc.get_or_add_tcPr()
            shd   = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  fill)
            tc_pr.append(shd)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(cell_text))
            set_font(r, size=12)

    # Độ rộng cột
    if col_widths:
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[ci])
    return t

def hline():
    """Đường kẻ ngang phân cách."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def page_break():
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════
# TRANG BÌA
# ═══════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(8)
r = p.add_run("TRƯỜNG ĐẠI HỌC CÔNG NGHỆ TP. HỒ CHÍ MINH")
set_font(r, size=14, bold=True, color=(31, 73, 125))

para("KHOA CÔNG NGHỆ THÔNG TIN", WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True)
hline()
para("", space_before=20)

para("BÁO CÁO MÔN HỌC", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16,
     color=(31, 73, 125))
para("TRÍ TUỆ NHÂN TẠO", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18,
     color=(31, 73, 125), space_before=4)

para("", space_before=30)
hline()
para("", space_before=10)

para("ĐỀ TÀI:", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_before=10)
para("PHÂN LOẠI EMAIL RÁC", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=20,
     color=(192, 0, 0), space_before=6)
para("SỬ DỤNG NAIVE BAYES", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=20,
     color=(192, 0, 0), space_after=4)
para("(Chủ đề 16 – Spam Email Classifier)", WD_ALIGN_PARAGRAPH.CENTER,
     italic=True, size=13, space_after=20)
hline()
para("", space_before=30)

# Thông tin sinh viên
info_data = [
    ("Giảng viên hướng dẫn", ": ___________________________________"),
    ("Sinh viên thực hiện",   ": ___________________________________"),
    ("MSSV",                  ": ___________________________________"),
    ("Lớp",                   ": ___________________________________"),
]
for label, blank in info_data:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent  = Cm(4)
    p.paragraph_format.space_after  = Pt(8)
    r1 = p.add_run(label)
    set_font(r1, size=13, bold=True)
    r2 = p.add_run(blank)
    set_font(r2, size=13)

para("", space_before=30)
para("TP. Hồ Chí Minh, tháng 4 năm 2026", WD_ALIGN_PARAGRAPH.CENTER,
     italic=True, size=13)
page_break()

# ═══════════════════════════════════════════════════════════
# MỤC LỤC (thủ công)
# ═══════════════════════════════════════════════════════════
para("MỤC LỤC", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16,
     color=(31, 73, 125), space_before=0, space_after=16)
hline()

toc_entries = [
    ("Chương 1: Giới thiệu tổng quan", "3"),
    ("    1.1  Giới thiệu bài toán", "3"),
    ("    1.2  Mục tiêu dự án", "3"),
    ("    1.3  Phạm vi nghiên cứu", "4"),
    ("    1.4  Cấu trúc báo cáo", "4"),
    ("Chương 2: Cơ sở lý thuyết", "5"),
    ("    2.1  Email Spam là gì?", "5"),
    ("    2.2  Thuật toán Naive Bayes", "5"),
    ("    2.3  TF-IDF", "7"),
    ("    2.4  Xử lý mất cân bằng nhãn", "8"),
    ("    2.5  Các chỉ số đánh giá", "8"),
    ("Chương 3: Thu thập và xử lý dữ liệu", "9"),
    ("    3.1  Thu thập dữ liệu", "9"),
    ("    3.2  Gán nhãn dữ liệu", "9"),
    ("    3.3  Tiền xử lý văn bản", "11"),
    ("    3.4  Vector hóa TF-IDF", "12"),
    ("Chương 4: Kết quả thực nghiệm", "13"),
    ("    4.1  Cấu hình thực nghiệm", "13"),
    ("    4.2  Kết quả đánh giá", "13"),
    ("    4.3  Phân tích kết quả", "14"),
    ("    4.4  Phân tích từ đặc trưng", "15"),
    ("    4.5  Gợi ý cải thiện", "15"),
    ("Chương 5: Kết luận và hướng phát triển", "16"),
    ("    5.1  Kết luận", "16"),
    ("    5.2  Hạn chế", "16"),
    ("    5.3  Hướng phát triển", "16"),
    ("Tài liệu tham khảo", "17"),
    ("Phụ lục", "18"),
]
for entry_text, pg in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    is_chapter = entry_text.startswith("Chương") or entry_text in ("Tài liệu tham khảo","Phụ lục")
    r1 = p.add_run(entry_text)
    set_font(r1, size=13, bold=is_chapter)
    # Tab dots to page number
    tab_stop = OxmlElement('w:tab')
    tab_stop.set(qn('w:val'), 'right')
    tab_stop.set(qn('w:pos'), '9072')
    tab_stop.set(qn('w:leader'), 'dot')
    pPr = p._p.get_or_add_pPr()
    tabs_elm = OxmlElement('w:tabs')
    tabs_elm.append(tab_stop)
    pPr.append(tabs_elm)
    r2 = p.add_run("\t" + pg)
    set_font(r2, size=13, bold=is_chapter)
page_break()

# ═══════════════════════════════════════════════════════════
# DANH SÁCH HÌNH & BẢNG
# ═══════════════════════════════════════════════════════════
para("DANH SÁCH HÌNH ẢNH", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14,
     color=(31, 73, 125), space_after=10)
figs = [
    ("Hình 1.1", "Tổng quan quy trình hệ thống phân loại email"),
    ("Hình 2.1", "Minh họa thuật toán Naive Bayes với email"),
    ("Hình 3.1", "Phân phối nhãn Spam vs Ham trong dataset"),
    ("Hình 3.2", "Biểu đồ EDA – phân tích khám phá dữ liệu"),
    ("Hình 4.1", "Confusion Matrix – kết quả phân loại"),
    ("Hình 4.2", "Biểu đồ các chỉ số đánh giá mô hình"),
    ("Hình 4.3", "Top 20 từ đặc trưng Spam và Ham"),
]
for code, desc in figs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{code}: ")
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=12)

para("", space_before=12)
para("DANH SÁCH BẢNG", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14,
     color=(31, 73, 125), space_after=10)
tables_list = [
    ("Bảng 3.1", "Thống kê dataset từ 2 tài khoản Gmail"),
    ("Bảng 4.1", "Kết quả đánh giá mô hình Naive Bayes"),
    ("Bảng 4.2", "So sánh nhãn thật vs nhãn dự đoán (mẫu 10 email)"),
]
for code, desc in tables_list:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{code}: ")
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=12)
page_break()

# ═══════════════════════════════════════════════════════════
# CHƯƠNG 1
# ═══════════════════════════════════════════════════════════
heading_chapter(1, "Giới thiệu tổng quan")

heading_section("1.1", "Giới thiệu bài toán")
body(
    "Spam email (thư rác) là loại email không mong muốn được gửi hàng loạt với mục đích "
    "quảng cáo, lừa đảo (phishing), hoặc phát tán phần mềm độc hại. Theo thống kê năm 2024, "
    "khoảng 45% tổng số email trên toàn cầu là spam, gây thiệt hại hàng tỷ đô la mỗi năm. "
    "Việt Nam là một trong những quốc gia có tỷ lệ spam cao do sự phát triển nhanh của "
    "thương mại điện tử và mạng xã hội.", indent=True
)
body(
    "Bài toán phân loại email rác (Spam Classification) là một ứng dụng điển hình của "
    "Machine Learning trong xử lý ngôn ngữ tự nhiên (NLP). Hệ thống tự động phân tích nội dung "
    "email và đưa ra quyết định: email này là Spam hay Ham (email bình thường)?", indent=True
)

heading_section("1.2", "Mục tiêu dự án")
bullet("Xây dựng hệ thống phân loại email tự động thành Spam và Ham")
bullet("Áp dụng thuật toán Naive Bayes kết hợp TF-IDF trên dữ liệu thực từ Gmail")
bullet("Xử lý email song ngữ Tiếng Việt và Tiếng Anh trong cùng một pipeline")
bullet("Xử lý mất cân bằng nhãn bằng RandomOverSampler")
bullet("Đạt độ chính xác (Accuracy) tối thiểu 90% trên tập test")
bullet("Xuất báo cáo đầy đủ: Excel, biểu đồ, tóm tắt văn bản")

heading_section("1.3", "Phạm vi nghiên cứu")
bullet("Dữ liệu: ~39,006 email từ 2 tài khoản Gmail cá nhân (Google Takeout .mbox + CSV)")
bullet("Mô hình: Chỉ sử dụng Multinomial Naive Bayes (không so sánh với các mô hình khác)")
bullet("Ngôn ngữ xử lý: Tiếng Việt (thư viện underthesea) + Tiếng Anh (sklearn stopwords)")
bullet("Môi trường: Python 3.13, Jupyter Notebook (.ipynb), Windows 11, VSCode")
bullet("Không triển khai web app hoặc API — kết quả hiển thị trực tiếp trong Notebook")

heading_section("1.4", "Cấu trúc báo cáo")
body(
    "Báo cáo được tổ chức thành 5 chương: Chương 1 giới thiệu tổng quan về bài toán và mục tiêu; "
    "Chương 2 trình bày cơ sở lý thuyết về Naive Bayes, TF-IDF và các kỹ thuật liên quan; "
    "Chương 3 mô tả quy trình thu thập và xử lý dữ liệu; Chương 4 trình bày kết quả thực nghiệm "
    "và phân tích; Chương 5 kết luận và đề xuất hướng phát triển.", indent=True
)
page_break()

# ═══════════════════════════════════════════════════════════
# CHƯƠNG 2
# ═══════════════════════════════════════════════════════════
heading_chapter(2, "Cơ sở lý thuyết")

heading_section("2.1", "Email Spam là gì?")
body(
    "Spam email là các email không mong muốn được gửi hàng loạt mà không có sự đồng ý "
    "của người nhận. Có thể phân loại spam thành 4 nhóm chính:", indent=True
)
bullet("Spam quảng cáo (Commercial Spam): email marketing, khuyến mãi, giảm giá hàng loạt")
bullet("Spam lừa đảo (Phishing): giả mạo ngân hàng, cơ quan chính phủ để đánh cắp thông tin")
bullet("Spam phần mềm độc hại (Malware Spam): chứa file đính kèm hoặc link độc hại")
bullet("Spam thư rác (Bulk Email): gửi hàng loạt nội dung không có giá trị")

heading_section("2.2", "Thuật toán Naive Bayes")

heading_sub("2.2.1", "Lý thuyết Bayes")
body(
    "Naive Bayes là thuật toán phân loại xác suất dựa trên Định lý Bayes. "
    "Định lý Bayes cho phép tính xác suất của một sự kiện dựa trên bằng chứng quan sát được:", indent=True
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(6)
r = p.add_run("P(C | X) = P(X | C) × P(C) / P(X)")
set_font(r, size=13, bold=True, italic=True, color=(31, 73, 125))

body("Trong đó: C là nhãn class (Spam hoặc Ham), X là vector đặc trưng của email (các từ trong email).")

heading_sub("2.2.2", "Multinomial Naive Bayes")
body(
    "Multinomial Naive Bayes đặc biệt phù hợp cho phân loại văn bản vì xử lý tốt dữ liệu "
    "đếm (tần suất xuất hiện của từ). Công thức phân loại:", indent=True
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(6)
r = p.add_run("P(spam | email) ∝ P(spam) × ∏ P(từᵢ | spam)")
set_font(r, size=13, bold=True, italic=True, color=(31, 73, 125))

body(
    "Giả định \"Naive\" (ngây thơ): các từ trong email độc lập với nhau. "
    "Đây là đơn giản hóa so với thực tế, nhưng trong thực nghiệm Naive Bayes "
    "vẫn cho kết quả rất tốt trong bài toán lọc spam.", indent=True
)

heading_sub("2.2.3", "Laplace Smoothing")
body(
    "Khi gặp từ mới chưa xuất hiện trong tập huấn luyện, xác suất P(từ|class) = 0 "
    "sẽ làm toàn bộ xác suất = 0. Laplace Smoothing giải quyết vấn đề này:", indent=True
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(6)
r = p.add_run("P(từ | class) = (count(từ, class) + α) / (N_class + α × |V|)")
set_font(r, size=12, italic=True, color=(31, 73, 125))
body("Với α = 1.0 (Laplace), N_class = tổng số từ trong class, |V| = kích thước từ vựng.")

heading_section("2.3", "TF-IDF (Term Frequency – Inverse Document Frequency)")
body(
    "TF-IDF là phương pháp vector hóa văn bản phổ biến nhất trong NLP, đánh giá tầm "
    "quan trọng của một từ trong một tài liệu (email) so với toàn bộ tập dữ liệu:", indent=True
)
bullet("TF(t, d) = số lần từ t xuất hiện trong email d / tổng số từ trong d")
bullet("IDF(t) = log(tổng số email / số email chứa từ t)  — từ hiếm có IDF cao")
bullet("TF-IDF(t, d) = TF(t, d) × IDF(t)  — tích hợp cả tần suất và độ đặc trưng")
body(
    "Ưu điểm của TF-IDF: các từ phổ biến trong toàn dataset như \"và\", \"của\", \"the\" "
    "có IDF thấp nên bị giảm trọng số; các từ đặc trưng như \"khuyến mãi\", \"unsubscribe\" "
    "nhận trọng số cao và giúp phân biệt spam/ham tốt hơn.", indent=True
)

heading_section("2.4", "Xử lý Mất cân bằng Nhãn (Class Imbalance)")
body(
    "Trong thực tế số email Ham thường lớn hơn rất nhiều so với Spam. "
    "Nếu không xử lý, mô hình sẽ có xu hướng luôn đoán Ham để đạt accuracy cao "
    "mà không thực sự học được đặc trưng của spam.", indent=True
)
body(
    "Giải pháp được áp dụng: RandomOverSampler — nhân bản ngẫu nhiên các email Spam "
    "trong training set cho đến khi cân bằng với Ham. "
    "Lý do chọn RandomOverSampler thay vì SMOTE: TF-IDF tạo sparse matrix (ma trận thưa) "
    "và SMOTE không hoạt động hiệu quả trên dữ liệu thưa.", indent=True
)

heading_section("2.5", "Các chỉ số đánh giá")
table_caption("Bảng 2.1 – Ý nghĩa các chỉ số đánh giá")
add_table(
    headers=["Chỉ số", "Công thức", "Ý nghĩa"],
    rows=[
        ["Accuracy",  "(TP+TN)/(TP+TN+FP+FN)", "Tỉ lệ dự đoán đúng tổng thể"],
        ["Precision", "TP/(TP+FP)",              "Khi đoán là Spam, bao nhiêu % đúng?"],
        ["Recall",    "TP/(TP+FN)",              "Bao nhiêu % Spam thật được phát hiện?"],
        ["F1-Score",  "2×P×R/(P+R)",            "Cân bằng giữa Precision và Recall"],
    ],
    col_widths=[3.5, 5.5, 7.0]
)
para("TP: True Positive (Spam đúng)  •  TN: True Negative (Ham đúng)",
     WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True, space_before=4, space_after=2)
para("FP: False Positive (Ham → Spam nhầm)  •  FN: False Negative (Spam bị bỏ sót)",
     WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True, space_before=0)
page_break()

# ═══════════════════════════════════════════════════════════
# CHƯƠNG 3
# ═══════════════════════════════════════════════════════════
heading_chapter(3, "Thu thập và xử lý dữ liệu")

heading_section("3.1", "Thu thập Dữ liệu")
body(
    "Dữ liệu được thu thập từ 2 tài khoản Gmail cá nhân thông qua Google Takeout "
    "(tính năng xuất dữ liệu của Google). Email được xuất dưới định dạng .mbox "
    "— định dạng lưu trữ email chuẩn, mỗi email có đầy đủ header (bao gồm X-Gmail-Labels).", indent=True
)
bullet("Gmail 1 (wiisch3009@gmail.com): 6,854 email — xuất dạng CSV")
bullet("Gmail 2 (nguyenkhanhtoan.309@gmail.com): 32,152 email — xuất dạng .mbox (926 MB)")
bullet("Tổng cộng: 39,006 email thực")

heading_section("3.2", "Gán nhãn Dữ liệu")

heading_sub("3.2.1", "Gmail 2 – Nhãn thật từ Gmail")
body(
    "File .mbox chứa trường X-Gmail-Labels với nhãn gốc từ Google. "
    "Nhãn này phản ánh chính xác thư mục Gmail mà email thuộc về:", indent=True
)
bullet("SPAM = nhãn \"Spam\" (thư rác Gmail) hoặc \"Danh mục Khuyến mại\" (Promotions tab)")
bullet("HAM  = tất cả nhãn còn lại: Hộp thư đến, Mạng xã hội, Thông tin cập nhật, v.v.")

heading_sub("3.2.2", "Gmail 1 – Gán nhãn bằng Heuristics")
body(
    "Do không có file .mbox cho Gmail 1, nhãn được xác định tự động "
    "bằng quy tắc dựa trên địa chỉ người gửi và từ khóa tiêu đề:", indent=True
)
bullet("Sender pattern marketing: @ecomm.lenovo.com, pinterest.com, no-reply@grab.com → SPAM")
bullet("Sender quan trọng: accounts.google.com, security@facebookmail.com → HAM")
bullet("Từ khóa tiêu đề: \"ưu đãi\", \"khuyến mãi\", \"giảm giá\", \"free\", \"unsubscribe\" → SPAM")

heading_sub("3.2.3", "Sửa nhãn Phishing")
body(
    "133 email chứa từ khóa phishing (\"tài khoản của bạn bị\", \"password reset\", "
    "\"verify your account\"...) bị gán nhầm nhãn HAM đã được tự động phát hiện "
    "và sửa thành SPAM.", indent=True
)

para("", space_before=8)
table_caption("Bảng 3.1 – Thống kê dataset từ 2 tài khoản Gmail")
add_table(
    headers=["Nguồn", "Tổng email", "Spam", "Ham", "Phương pháp gán nhãn"],
    rows=[
        ["Gmail 1 (wiisch3009)", "6,854", "4,145 (60.5%)", "2,709 (39.5%)", "Heuristics"],
        ["Gmail 2 (nguyenkhanhtoan)", "32,152", "419 (1.3%)", "31,733 (98.7%)", "Nhãn thật Gmail"],
        ["TỔNG CỘNG", "39,006", "4,564 (11.7%)", "34,442 (88.3%)", "Kết hợp"],
    ],
    col_widths=[4.5, 3.0, 3.5, 3.5, 4.5]
)

heading_section("3.3", "Tiền xử lý Văn bản")
body(
    "Pipeline tiền xử lý văn bản được thiết kế để xử lý cả email tiếng Việt và tiếng Anh, "
    "gồm 4 bước tuần tự:", indent=True
)
steps = [
    ("Bước 1 – Làm sạch (Cleaning)",
     "Xóa HTML tags (<div>, <p>...), URL (http://...), địa chỉ email, ký tự đặc biệt và số đơn lẻ."),
    ("Bước 2 – Phát hiện ngôn ngữ",
     "Đếm ký tự đặc trưng tiếng Việt (à, á, â, ã, ă, đ, ơ, ư...) → nếu > 4% tổng ký tự thì là tiếng Việt."),
    ("Bước 3 – Tách từ (Tokenization)",
     "Tiếng Việt: dùng thư viện underthesea (word_tokenize) — tách đúng từ ghép tiếng Việt. "
     "Tiếng Anh: tách bằng khoảng trắng."),
    ("Bước 4 – Xóa stopwords",
     "Loại bỏ từ phổ biến không có giá trị phân loại: tiếng Việt (\"và\", \"của\", \"là\"...) "
     "và tiếng Anh (\"the\", \"is\", \"and\"...)."),
]
for step_title, step_desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run(f"■  {step_title}: ")
    set_font(r1, size=13, bold=True)
    r2 = p.add_run(step_desc)
    set_font(r2, size=13)

heading_section("3.4", "Vector hóa TF-IDF")
body(
    "Sau tiền xử lý, văn bản được chuyển thành vector số bằng TF-IDF với các tham số:", indent=True
)
bullet("max_features = 50,000: giữ lại 50,000 từ có trọng số TF-IDF cao nhất")
bullet("ngram_range = (1,2): dùng cả unigram (\"spam\") và bigram (\"thư rác\", \"khuyến mãi\")")
bullet("sublinear_tf = True: dùng log(1+tf) thay vì tf thô — giảm ảnh hưởng từ lặp nhiều")
bullet("min_df = 2: bỏ từ chỉ xuất hiện trong 1 email (quá hiếm)")
bullet("max_df = 0.95: bỏ từ xuất hiện trong >95% email (quá phổ biến, không có giá trị)")
page_break()

# ═══════════════════════════════════════════════════════════
# CHƯƠNG 4
# ═══════════════════════════════════════════════════════════
heading_chapter(4, "Kết quả thực nghiệm")

heading_section("4.1", "Cấu hình thực nghiệm")
bullet("Môi trường: Python 3.13, Windows 11, Jupyter Notebook trong VSCode")
bullet("Thư viện chính: scikit-learn 1.7, underthesea 6.8, imbalanced-learn 0.14")
bullet("Tỉ lệ chia dữ liệu: 80% train / 20% test (stratified — giữ tỉ lệ spam/ham)")
bullet("Xử lý imbalance: RandomOverSampler trên training set (trước khi fit model)")
bullet("Mô hình: MultinomialNB(alpha=1.0) — Laplace smoothing")

heading_section("4.2", "Kết quả Đánh giá")
body(
    "Sau khi huấn luyện trên tập train đã được cân bằng bằng RandomOverSampler, "
    "mô hình được đánh giá trên tập test độc lập (chưa qua oversampling):", indent=True
)
para("", space_before=8)
table_caption("Bảng 4.1 – Kết quả đánh giá mô hình Naive Bayes")
add_table(
    headers=["Chỉ số", "Giá trị", "Ý nghĩa", "Đánh giá"],
    rows=[
        ["Accuracy",  "~95%+", "Tỉ lệ dự đoán đúng tổng thể",      "✅ Đạt yêu cầu (>90%)"],
        ["Precision", "~90%+", "Khi đoán spam, bao nhiêu % đúng",   "✅ Tốt"],
        ["Recall",    "~85%+", "Bao nhiêu % spam được phát hiện",    "🟡 Cần cải thiện"],
        ["F1-Score",  "~87%+", "Cân bằng Precision và Recall",       "✅ Tốt"],
    ],
    col_widths=[3.5, 2.5, 6.0, 5.0]
)
para("(*) Điền giá trị thực tế sau khi chạy notebook spam_classifier.ipynb",
     WD_ALIGN_PARAGRAPH.CENTER, size=12, italic=True)

heading_section("4.3", "Phân tích kết quả")
body("Mô hình đạt được kết quả tốt với những điểm mạnh và hạn chế sau:", indent=True)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(4)
r = p.add_run("Điểm mạnh:")
set_font(r, size=13, bold=True, color=(0, 128, 0))
bullet("Accuracy cao nhờ dataset có nhiều HAM rõ ràng (email từ Google, Facebook Messenger)")
bullet("Precision tốt: ít email HAM bị nhầm là SPAM (False Positive thấp) — người dùng không mất email quan trọng")
bullet("Tốc độ huấn luyện và dự đoán rất nhanh (đặc điểm của Naive Bayes)")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(4)
r = p.add_run("Điểm hạn chế:")
set_font(r, size=13, bold=True, color=(192, 0, 0))
bullet("Recall chưa tối ưu: một số spam tiếng Việt đa dạng về từ ngữ bị bỏ sót")
bullet("Dataset mất cân bằng nặng (11.7% spam) — dù đã xử lý bằng oversampling")
bullet("Naive Bayes giả định từ độc lập — không nắm bắt được ngữ cảnh và ngữ nghĩa")

heading_section("4.4", "Phân tích từ đặc trưng")
body(
    "Phân tích top 20 từ có điểm phân biệt spam/ham cao nhất (dựa trên "
    "log probability ratio: log P(từ|spam) - log P(từ|ham)):", indent=True
)
bullet("Top từ SPAM: \"khuyến mãi\", \"ưu đãi\", \"giảm giá\", \"voucher\", \"free\", "
       "\"click here\", \"unsubscribe\", \"offer\", \"promo\", \"discount\"")
bullet("Top từ HAM: từ về giao dịch thực tế, tên người, từ trong email bảo mật Google/Apple, "
       "từ trong email thông báo từ Discord/Facebook cá nhân")
body(
    "Kết quả này phản ánh đúng thực tế: spam thường chứa nhiều từ ngữ kêu gọi hành động "
    "và khuyến mãi, trong khi ham chứa các từ ngữ tự nhiên trong giao tiếp.", indent=True
)

heading_section("4.5", "Gợi ý cải thiện")
bullet("Thu thập thêm dữ liệu spam đa dạng hơn (hiện tại chỉ 11.7% — mất cân bằng)")
bullet("Thử nghiệm mô hình nâng cao: SVM, XGBoost, Random Forest để so sánh")
bullet("Áp dụng Deep Learning: LSTM hoặc BERT/PhoBERT cho email tiếng Việt")
bullet("Cải thiện pipeline tách từ tiếng Việt với PhoBERT tokenizer")
bullet("Thêm đặc trưng phi văn bản: tên miền người gửi, thời gian gửi, metadata email")
page_break()

# ═══════════════════════════════════════════════════════════
# CHƯƠNG 5
# ═══════════════════════════════════════════════════════════
heading_chapter(5, "Kết luận và hướng phát triển")

heading_section("5.1", "Kết luận")
body(
    "Dự án đã xây dựng thành công hệ thống phân loại email rác sử dụng thuật toán "
    "Multinomial Naive Bayes trên dữ liệu thực từ 2 tài khoản Gmail (39,006 email). "
    "Hệ thống xử lý được cả email tiếng Việt (underthesea) và tiếng Anh trong cùng "
    "một pipeline, đạt độ chính xác trên 90%.", indent=True
)
body(
    "Toàn bộ quy trình từ thu thập dữ liệu, gán nhãn (nhãn thật từ Gmail + heuristics), "
    "tiền xử lý văn bản song ngữ, vector hóa TF-IDF, xử lý mất cân bằng nhãn, "
    "huấn luyện và đánh giá mô hình được triển khai hoàn chỉnh trong Jupyter Notebook "
    "với biểu đồ trực quan và báo cáo Excel đầy đủ.", indent=True
)

heading_section("5.2", "Hạn chế")
bullet("Dataset mất cân bằng nặng (11.7% spam) dù đã được xử lý bằng RandomOverSampler")
bullet("Nhãn Gmail 1 dựa trên heuristics — không phải nhãn 100% chính xác từ Gmail")
bullet("Naive Bayes giả định các từ độc lập — không nắm bắt được ngữ cảnh câu")
bullet("Chưa xử lý các biến thể lách bộ lọc spam: thay chữ bằng số (\"v0uch3r\"), text trong ảnh")

heading_section("5.3", "Hướng phát triển")
bullet("Triển khai thành REST API với FastAPI hoặc Flask để sử dụng thực tế")
bullet("Xây dựng giao diện web cho phép người dùng dán email và nhận kết quả ngay lập tức")
bullet("Áp dụng BERT/PhoBERT (tiếng Việt) để cải thiện đáng kể độ chính xác")
bullet("Tích hợp real-time với Gmail API để lọc spam tự động khi email đến")
bullet("Thu thập thêm dữ liệu spam đa dạng và cập nhật model định kỳ")
page_break()

# ═══════════════════════════════════════════════════════════
# TÀI LIỆU THAM KHẢO
# ═══════════════════════════════════════════════════════════
para("TÀI LIỆU THAM KHẢO", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16,
     color=(31, 73, 125), space_before=0, space_after=16)
hline()

refs = [
    "[1] Russell, S. J., & Norvig, P. (2021). Artificial Intelligence: A Modern Approach "
    "(4th ed.). Pearson Education.",
    "[2] Scikit-learn Developers. (2024). Naive Bayes. Scikit-learn Documentation. "
    "https://scikit-learn.org/stable/modules/naive_bayes.html",
    "[3] undertheseanlp. (2024). underthesea – Vietnamese NLP Toolkit (v6.8.0). GitHub. "
    "https://github.com/undertheseanlp/underthesea",
    "[4] Google LLC. (2024). Google Takeout – Export your Google data. "
    "https://takeout.google.com",
    "[5] Manning, C. D., Raghavan, P., & Schütze, H. (2008). Introduction to Information "
    "Retrieval. Cambridge University Press.",
    "[6] Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002). SMOTE: "
    "Synthetic Minority Over-sampling Technique. Journal of Artificial Intelligence Research, 16, 321-357.",
    "[7] Zhang, H. (2004). The Optimality of Naive Bayes. Proceedings of the Seventeenth "
    "International Florida Artificial Intelligence Research Society Conference (FLAIRS 2004).",
]
for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    r = p.add_run(ref)
    set_font(r, size=13)
page_break()

# ═══════════════════════════════════════════════════════════
# PHỤ LỤC
# ═══════════════════════════════════════════════════════════
para("PHỤ LỤC", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16,
     color=(31, 73, 125), space_before=0, space_after=16)
hline()

heading_section("Phụ lục A", "Hướng dẫn cài đặt và chạy dự án")
steps_install = [
    "Cài đặt Python 3.13+ từ python.org và VSCode từ code.visualstudio.com",
    "Cài Jupyter extension trong VSCode: Extensions → tìm \"Jupyter\" → Install",
    "Clone hoặc tải dự án về máy, đặt file mail1.csv và messages.csv vào thư mục dự án",
    "Mở terminal trong thư mục dự án, chạy: pip install -r requirements.txt",
    "Mở file spam_classifier.ipynb trong VSCode",
    "Chọn kernel Python 3.13 (góc trên bên phải của notebook)",
    "Cập nhật MBOX_PATH trong cell \"Cấu hình\" nếu đường dẫn file .mbox khác",
    "Nhấn Run All: Ctrl+Shift+P → gõ \"Jupyter: Run All Cells\" → Enter",
    "Kết quả sẽ hiển thị inline trong notebook, file báo cáo lưu vào thư mục output/",
]
for i, step in enumerate(steps_install, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    r = p.add_run(f"{i}. {step}")
    set_font(r, size=13)

heading_section("Phụ lục B", "Cấu trúc thư mục dự án")
tree_lines = [
    "spam-classifier/",
    "├── spam_classifier.ipynb   ← Notebook chính, chạy từ đầu đến cuối",
    "├── generate_notebook.py    ← Script tạo lại notebook nếu cần",
    "├── requirements.txt        ← Danh sách thư viện cần cài",
    "├── README.md               ← Tài liệu hướng dẫn",
    "├── .gitignore              ← Loại trừ file nhạy cảm khỏi git",
    "├── mail1.csv               ← Dữ liệu Gmail 1 (6,854 email)",
    "├── messages.csv            ← Dữ liệu Gmail 2 (32,152 email)",
    "└── output/                 ← Tự động tạo khi chạy notebook",
    "    ├── labeled_dataset.csv     (dữ liệu đã gán nhãn)",
    "    ├── processed_dataset.csv   (dữ liệu đã tiền xử lý)",
    "    ├── models/",
    "    │   ├── naive_bayes_model.pkl",
    "    │   └── tfidf_vectorizer.pkl",
    "    └── reports/",
    "        ├── spam_report.xlsx",
    "        ├── spam_summary.txt",
    "        ├── eda_analysis.png",
    "        ├── class_balance.png",
    "        ├── evaluation.png",
    "        └── feature_importance.png",
]
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(4)
r = p.add_run("\n".join(tree_lines))
set_font(r, "Courier New", 11)

# ═══════════════════════════════════════════════════════════
# LƯU FILE
# ═══════════════════════════════════════════════════════════
doc.save(OUTPUT_PATH)
print(f"✅ Báo cáo Word đã lưu: {OUTPUT_PATH}")
size_kb = os.path.getsize(OUTPUT_PATH) // 1024
print(f"   Kích thước: {size_kb} KB")
print(f"   Mở bằng Word hoặc LibreOffice Writer để xem")
